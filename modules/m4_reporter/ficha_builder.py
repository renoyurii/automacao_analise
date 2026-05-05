"""
Gerador da Ficha de Verificação de Segurança da Informação — DESEG/TJRJ.

Replica o layout oficial produzido manualmente, seção por seção.
Design: limpo, minimalista, sem poluição visual. Status de conformidade
em destaque com código de cores; dados brutos em fonte monoespaçada.
"""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from config import REPORT_HEADER_LINE1, REPORT_HEADER_LINE2, REPORT_FOOTER
from .styles import (
    BLUE_HD, DARK, FONT_BODY, FONT_GRADE, FONT_MONO, FONT_NAME,
    FONT_NAME_MONO, FONT_SECTION, FONT_SMALL, FONT_TITLE, GRAY,
    GREEN, LIGHT, MID, ORANGE, RED, STATUS_COLOR, WHITE,
    remove_paragraph_spacing, set_cell_bg, set_cell_vertical_align,
    set_col_width, set_paragraph_spacing, set_table_borders,
)

_MARGIN = Cm(2.5)

# Rótulos dos protocolos TLS exibidos na ficha (ordem exata da ficha oficial)
_TLS_ROWS = [
    ("TLS 1.3", "TLS 1.3"),
    ("TLS 1.2", "TLS 1.2"),
    ("TLS 1.1", "TLS 1.1"),
    ("TLS 1.0", "TLS 1.0"),
    ("SSL 3.0", "SSL3 – SEGURANÇA"),
    ("SSL 2.0", "SSL2 – SEGURANÇA"),
]


# ── Interface pública ─────────────────────────────────────────────────────────

def build_ficha(result_data: dict[str, Any], output_path: str | Path) -> str:
    """
    Gera a Ficha de Verificação em .docx.

    Parâmetros:
        result_data — saída do M3 (evaluate())
        output_path — caminho de saída (criado se não existir)

    Retorna o caminho do arquivo gerado.
    """
    doc = Document()
    _setup_document(doc)

    _add_header(doc, result_data)
    _add_classificacao(doc, result_data)
    _add_disponibilidade(doc, result_data)
    _add_integridade(doc, result_data)
    _add_aplicacoes(doc, result_data)
    _add_hsts(doc, result_data)
    _add_criptografia(doc, result_data)
    _add_seguranca_rede(doc, result_data)
    _add_recomendacoes(doc)
    _add_conclusao(doc, result_data)
    _add_footer_text(doc)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)


# ── Configuração do documento ─────────────────────────────────────────────────

def _setup_document(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin    = _MARGIN
        section.bottom_margin = _MARGIN
        section.left_margin   = _MARGIN
        section.right_margin  = _MARGIN

    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_BODY
    style.font.color.rgb = DARK


# ── Seção 0: Cabeçalho institucional ─────────────────────────────────────────

def _add_header(doc: Document, rd: dict) -> None:
    p = doc.add_paragraph()
    _run(p, "Homologação de Leiloeiros e Corretores de Imóveis",
         bold=True, size=FONT_TITLE, color=BLUE_HD)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=4)

    p2 = doc.add_paragraph()
    _run(p2, "Lista de Verificação de Segurança da Informação",
         bold=True, size=FONT_SECTION, color=DARK)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p2, before=0, after=8)

    for text in (REPORT_HEADER_LINE1, REPORT_HEADER_LINE2):
        p = doc.add_paragraph()
        _run(p, text, size=FONT_SMALL, italic=True)
        set_paragraph_spacing(p, before=0, after=4)

    # URL do leiloeiro em destaque
    p = doc.add_paragraph()
    _run(p, rd.get("url", rd.get("domain", "")), bold=True, size=FONT_BODY, color=BLUE_HD)
    set_paragraph_spacing(p, before=6, after=6)

    doc.add_paragraph()  # espaço


# ── Seção 1: Classificação Geral SSL Labs ─────────────────────────────────────

def _add_classificacao(doc: Document, rd: dict) -> None:
    ssl = rd.get("raw", {}).get("ssl_labs") or {}
    grade = ssl.get("grade")
    scores = ssl.get("scores") or {}
    cert_valid = ssl.get("cert_valid")

    if not grade:
        return   # SSL Labs indisponível — omite a seção

    _section_label(doc, "CLASSIFICAÇÃO GERAL")

    # Grade em destaque (grande, centralizado)
    p_grade = doc.add_paragraph()
    r = p_grade.add_run(grade)
    r.font.name = FONT_NAME
    r.font.size = FONT_GRADE
    r.font.bold = True
    grade_color = GREEN if grade in ("A+", "A") else (ORANGE if grade == "B" else RED)
    r.font.color.rgb = grade_color
    p_grade.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p_grade, before=4, after=4)

    # Tabela de scores
    score_rows = [
        ("CERTIFICADO", 100 if cert_valid else (0 if cert_valid is False else "N/A")),
        ("SUPORTE DE PROTOCOLO", scores.get("suporte_protocolo", "N/A")),
        ("CHAVES",               scores.get("chaves", "N/A")),
        ("FORÇA DE CIFRA",       scores.get("forca_cifra", "N/A")),
    ]

    table = doc.add_table(rows=1, cols=2)
    set_table_borders(table)
    hdr = table.rows[0].cells
    _hdr_cell(hdr[0], "Critério")
    _hdr_cell(hdr[1], "Pontuação")

    for label, value in score_rows:
        row = table.add_row().cells
        _body_cell(row[0], label)
        _body_cell(row[1], str(value), align=WD_ALIGN_PARAGRAPH.CENTER)

    set_col_width(table, 0, 10.0)
    set_col_width(table, 1, 3.5)
    doc.add_paragraph()


# ── Seção 1: Disponibilidade ──────────────────────────────────────────────────

def _add_disponibilidade(doc: Document, rd: dict) -> None:
    raw_secs         = rd.get("raw", {}).get("claimed_raw_sections", {})
    llm_ev           = rd.get("raw", {}).get("llm_evidence", {}) or {}
    image_page_count = rd.get("raw", {}).get("image_page_count", 0) or 0
    _section_label(doc, "1. Disponibilidade")

    # (label, raw_section_key, llm_evidence_key)
    items = [
        ("Redundância de serviço",     "redundancia", "redundancy"),
        ("Backup e recuperação",       "backup",      "backup"),
        ("Recurso contínuo de energia","energia",     "energy"),
    ]

    table = doc.add_table(rows=1, cols=2)
    set_table_borders(table)
    hdr = table.rows[0].cells
    _hdr_cell(hdr[0], "Item")
    _hdr_cell(hdr[1], "Declarado pelo leiloeiro")

    for label, sec_key, llm_key in items:
        row = table.add_row().cells
        _body_cell(row[0], label)

        # Prioridade: LLM evidence (citação direta) > raw_sections (regex)
        llm_text = (llm_ev.get(llm_key, "") or "").strip()
        raw_text = raw_secs.get(sec_key, "").strip()
        evidence = llm_text if llm_text else raw_text
        _evidence_cell(row[1], evidence, image_page_count)

    set_col_width(table, 0, 5.5)
    set_col_width(table, 1, 11.0)
    doc.add_paragraph()


def _clean_evidence_text(text: str) -> str:
    """
    Remove ruídos típicos de PDFs em formato slide-deck/multi-coluna, em que
    cabeçalhos, rodapés e fragmentos do template da demanda se intercalam ao
    texto da resposta.
    """
    patterns = [
        # Cabeçalhos / rodapés do slide-deck
        r"Evidências?\s+(?:Brame\s+)?Leilões",
        r"Demandas?\s*",
        r"Despacho\s*[-–]\s*TJ/[\w/]+",
        r"Source:\s*\w+\s*\d*",
        # Templates de demanda (perguntas do TJ)
        r"Informar\s+a\s+exist[êe]ncia\s+de:?\s*●?\s*",
        r"●\s*redund[âa]ncia\s+de\s+servi[çc]os;?\s*",
        r"●\s*rotina\s+de\s+backup\s+e\s+recupera[çc][ãa]o;?\s*",
        r"●\s*recurso\s+cont[íi]nuo\s+de\s+energia;?\s*",
        # Outras demandas que vazam para o trecho
        r"Precisamos\s+que\s+sejam\s+informadas[^.]*\.",
        r"Solicitamos[^.]*\.",
        r"Para\s+o\s+SSL[^.]*\.",
        r"Al[ée]m\s+disso,\s+informar[^.]*\.",
        # Cabeçalho de norma interna repetido em cada página
        # (ex: "CÓDIGO N.026 NORMA VERSÃO V.001 PUBLICADO EM: 12/07/2024 VÁLIDO ATÉ: 12/07/2026")
        r"CÓDIGO\s+N\.\d+\s+NORMA\s+VERSÃO\s+V\.\d+[^\n]*",
        r"PUBLICADO\s+EM:\s+\d{2}/\d{2}/\d{4}[^\n]*",
        r"VÁLIDO\s+ATÉ:\s+\d{2}/\d{2}/\d{4}[^\n]*",
    ]
    cleaned = text
    for pat in patterns:
        cleaned = re.sub(pat, " ", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*●\s*●", " ●", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    cleaned = re.sub(r"^[;:.,●\s]+", "", cleaned)

    if cleaned and cleaned[0].islower():
        m = re.search(r"[.!?]\s+", cleaned[:160])
        if m:
            cleaned = cleaned[m.end():].lstrip()
    return cleaned


def _evidence_cell(cell, raw_text: str, image_page_count: int) -> None:
    """
    Apresenta a evidência da declaração de forma clara:
      - Inferência: rótulo destacado + motivo + trecho fonte entre aspas
      - Direta:     trecho citado entre aspas, em itálico
      - Vazio:      "Não informado" simples
    """
    # Limpa o parágrafo padrão da célula
    cell_p = cell.paragraphs[0]
    remove_paragraph_spacing(cell_p)
    set_cell_vertical_align(cell, "top")

    # Caso 1 — Inferência indireta
    if raw_text.startswith("[INFERIDO] "):
        body = raw_text[len("[INFERIDO] "):]
        # body tem o formato: "<motivo>. Trecho do documento: \"<quote>\""
        if 'Trecho do documento: "' in body:
            reason, quote = body.split('Trecho do documento: "', 1)
            quote = quote.rstrip('"').strip()
        else:
            reason, quote = body, ""
        reason = reason.rstrip(". ").strip()

        # Linha 1: rótulo "Inferido" em laranja + motivo
        r_label = cell_p.add_run("Inferido — ")
        r_label.font.name  = FONT_NAME
        r_label.font.size  = FONT_SMALL
        r_label.font.bold  = True
        r_label.font.color.rgb = ORANGE
        r_reason = cell_p.add_run(reason + ".")
        r_reason.font.name = FONT_NAME
        r_reason.font.size = FONT_SMALL
        r_reason.font.color.rgb = DARK

        if quote:
            quote = _clean_evidence_text(quote)
            p2 = cell.add_paragraph()
            remove_paragraph_spacing(p2)
            r_q_lbl = p2.add_run('Trecho do documento: ')
            r_q_lbl.font.name = FONT_NAME
            r_q_lbl.font.size = FONT_SMALL
            r_q_lbl.font.color.rgb = GRAY
            r_quote = p2.add_run('“' + quote + '”')
            r_quote.font.name = FONT_NAME
            r_quote.font.size = FONT_SMALL
            r_quote.font.italic = True
            r_quote.font.color.rgb = DARK
        return

    # Caso 2 — Declaração direta (mostra o trecho completo)
    if raw_text:
        _sep = ' [...] '
        parts = [
            _clean_evidence_text(p.strip())
            for p in raw_text.split(_sep)
            if p.strip()
        ]
        full_text = _sep.join(p for p in parts if p)
        r = cell_p.add_run('“' + full_text + '”')
        r.font.name = FONT_NAME
        r.font.size = FONT_SMALL
        r.font.italic = True
        r.font.color.rgb = DARK
        return

    # Caso 3 — Sem evidência
    r = cell_p.add_run("Não informado")
    r.font.name = FONT_NAME
    r.font.size = FONT_SMALL
    r.font.color.rgb = GRAY


# ── Seção 2: Integridade ──────────────────────────────────────────────────────

def _add_integridade(doc: Document, rd: dict) -> None:
    raw_block = rd.get("raw", {}).get("headers_raw_block", "")
    check = rd.get("checks", {}).get("integridade", {})

    _section_label(doc, "2. Integridade (Presença de firewall e/ou detecção de intrusão; Firewall, WAF, IPS/IDS)")

    # Status em linha
    status = check.get("status", "NÃO VERIFICÁVEL")
    p = doc.add_paragraph()
    _run(p, "Status: ", bold=True)
    _run(p, status, bold=True, color=STATUS_COLOR.get(status, DARK))
    set_paragraph_spacing(p, 0, 3)

    # Rótulo do bloco bruto
    lbl = doc.add_paragraph()
    _run(lbl, "INTEGRIDADE (PRESENÇA DE FIREWALL, WAF, BALANCEADOR, IPS/IDS)",
         bold=True, size=FONT_SMALL)
    set_paragraph_spacing(lbl, 4, 2)

    # Bloco de headers HTTP em fonte monoespaçada (máx. 50 linhas para não inflar o doc)
    if raw_block:
        lines = raw_block.splitlines()[:50]
        p_raw = doc.add_paragraph()
        p_raw.paragraph_format.left_indent = Cm(0.5)
        for i, line in enumerate(lines):
            r = p_raw.add_run(line)
            r.font.name = FONT_NAME_MONO
            r.font.size = FONT_MONO
            if i < len(lines) - 1:
                p_raw.add_run("\n")
        set_paragraph_spacing(p_raw, 0, 6)

    doc.add_paragraph()


# ── Seção 3: Aplicações Atualizadas ──────────────────────────────────────────

def _add_aplicacoes(doc: Document, rd: dict) -> None:
    technologies = rd.get("raw", {}).get("technologies", [])
    check = rd.get("checks", {}).get("aplicacoes", {})

    _section_label(doc, "3. Aplicações Atualizadas")

    # Status geral EOL
    status = check.get("status", "NÃO VERIFICÁVEL")
    p = doc.add_paragraph()
    _run(p, "Status EOL: ", bold=True)
    _run(p, status, bold=True, color=STATUS_COLOR.get(status, DARK))
    set_paragraph_spacing(p, 0, 4)

    if not technologies:
        _body_para(doc, "Nenhuma tecnologia detectada.")
        doc.add_paragraph()
        return

    lbl = doc.add_paragraph()
    _run(lbl, "APLICAÇÕES", bold=True, size=FONT_SMALL)
    set_paragraph_spacing(lbl, 2, 2)

    table = doc.add_table(rows=1, cols=4)
    set_table_borders(table)
    hdr = table.rows[0].cells
    _hdr_cell(hdr[0], "Categoria")
    _hdr_cell(hdr[1], "Tecnologia")
    _hdr_cell(hdr[2], "Versão")
    _hdr_cell(hdr[3], "EOL")

    for i, tech in enumerate(technologies):
        row = table.add_row().cells
        bg = None if i % 2 == 0 else "F5F5F5"
        for cell, txt in zip(row, [
            tech.get("category", ""),
            tech.get("name", ""),
            tech.get("version") or "N/A",
        ]):
            _body_cell(cell, txt)
            if bg:
                set_cell_bg(cell, bg)

        eol = tech.get("eol")
        eol_txt = "EOL" if eol is True else ("OK" if eol is False else "—")
        eol_color = "C62828" if eol is True else ("2E7D32" if eol is False else "546E7A")
        p_eol = row[3].paragraphs[0]
        r = p_eol.add_run(eol_txt)
        r.font.name = FONT_NAME
        r.font.size = FONT_SMALL
        r.font.bold = eol is True
        r.font.color.rgb = RGBColor.from_string(eol_color)
        p_eol.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if bg:
            set_cell_bg(row[3], bg)

    set_col_width(table, 0, 4.0)
    set_col_width(table, 1, 4.5)
    set_col_width(table, 2, 2.5)
    set_col_width(table, 3, 1.5)
    doc.add_paragraph()


# ── Seção 4: HSTS ─────────────────────────────────────────────────────────────

def _add_hsts(doc: Document, rd: dict) -> None:
    check = rd.get("checks", {}).get("hsts", {})
    found = check.get("found")
    status = check.get("status", "NÃO VERIFICÁVEL")

    _section_label(doc, "4. HSTS")

    table = doc.add_table(rows=1, cols=2)
    set_table_borders(table)
    hdr = table.rows[0].cells
    _hdr_cell(hdr[0], "HSTS")
    _hdr_cell(hdr[1], "Status")

    row = table.add_row().cells
    _body_cell(row[0], "Segurança Estrita de Transporte (HSTS)")

    # Exibe SIM/NÃO com cor, ou o status completo se NÃO VERIFICÁVEL
    if status == "NÃO VERIFICÁVEL":
        _status_cell(row[1], status)
    else:
        sim_nao = "SIM" if found else "NÃO"
        p = row[1].paragraphs[0]
        r = p.add_run(sim_nao)
        r.font.name = FONT_NAME
        r.font.size = FONT_BODY
        r.font.bold = True
        r.font.color.rgb = GREEN if found else RED
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Nota sobre max-age se HSTS presente mas subótimo
    detail = check.get("detail", "")
    if "max-age" in detail and status == "ATENÇÃO":
        p_note = doc.add_paragraph()
        _run(p_note, f"  ⚠ {detail}", size=FONT_SMALL, color=ORANGE)
        set_paragraph_spacing(p_note, 2, 2)

    set_col_width(table, 0, 10.0)
    set_col_width(table, 1, 3.5)
    doc.add_paragraph()


# ── Seção 5: Criptografia TLS ─────────────────────────────────────────────────

def _add_criptografia(doc: Document, rd: dict) -> None:
    cripto = rd.get("checks", {}).get("criptografia", {})

    _section_label(doc, "5. Criptografia — SSL — TLS 1.2 | 1.3")

    lbl = doc.add_paragraph()
    _run(lbl, "CRIPTOGRAFIA - SSL - TLS 1.2|1.3", bold=True, size=FONT_SMALL)
    set_paragraph_spacing(lbl, 2, 2)

    table = doc.add_table(rows=1, cols=2)
    set_table_borders(table)
    hdr = table.rows[0].cells
    _hdr_cell(hdr[0], "Protocolo")
    _hdr_cell(hdr[1], "Possui?")

    for key, display_name in _TLS_ROWS:
        check = cripto.get(key, {})
        found = check.get("found")
        status = check.get("status", "NÃO VERIFICÁVEL")

        row = table.add_row().cells
        _body_cell(row[0], display_name)

        if status == "NÃO VERIFICÁVEL":
            _status_cell(row[1], "NÃO VERIFICÁVEL")
        else:
            # Para protocolos que devem estar DESABILITADOS (TLS 1.1, 1.0, SSL*)
            # "NÃO" (found=False) é o estado CORRETO → verde
            # "SIM" (found=True) é o estado INCORRETO → vermelho
            should_be_disabled = key in ("TLS 1.1", "TLS 1.0", "SSL 3.0", "SSL 2.0")
            if should_be_disabled:
                sim_nao = "NÃO" if not found else "SIM"
                color = GREEN if not found else RED
            else:
                sim_nao = "SIM" if found else "NÃO"
                color = GREEN if found else RED

            p = row[1].paragraphs[0]
            r = p.add_run(sim_nao)
            r.font.name = FONT_NAME
            r.font.size = FONT_BODY
            r.font.bold = True
            r.font.color.rgb = color
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_col_width(table, 0, 10.0)
    set_col_width(table, 1, 3.5)
    doc.add_paragraph()


# ── Seção 6: Segurança de Rede (WHOIS) ───────────────────────────────────────

def _add_seguranca_rede(doc: Document, rd: dict) -> None:
    whois_raw = rd.get("raw", {}).get("whois_raw", "")
    check = rd.get("checks", {}).get("seguranca_rede", {})

    _section_label(doc, "6. Segurança de Rede")

    lbl = doc.add_paragraph()
    _run(lbl, "INFORMAÇÕES DE REDE", bold=True, size=FONT_SMALL)
    set_paragraph_spacing(lbl, 2, 2)

    # Filtra e limita o WHOIS bruto (remove linhas de comentário #)
    if whois_raw:
        lines = [
            ln for ln in whois_raw.splitlines()
            if ln.strip() and not ln.strip().startswith("%")
            and not ln.strip().startswith("#")
        ][:40]
        p_raw = doc.add_paragraph()
        p_raw.paragraph_format.left_indent = Cm(0.5)
        for i, line in enumerate(lines):
            r = p_raw.add_run(line)
            r.font.name = FONT_NAME_MONO
            r.font.size = FONT_MONO
            if i < len(lines) - 1:
                p_raw.add_run("\n")
        set_paragraph_spacing(p_raw, 0, 6)

    # Alerta de domínio expirando
    if check.get("status") == "NÃO CONFORME":
        p_alert = doc.add_paragraph()
        _run(p_alert, f"⚠ {check.get('detail', '')}", bold=True, color=RED)
        set_paragraph_spacing(p_alert, 2, 4)

    doc.add_paragraph()


# ── Recomendações (texto fixo do template DESEG) ──────────────────────────────

def _add_recomendacoes(doc: Document) -> None:
    _section_label(doc, "Recomendações")

    recomendacoes = [
        (
            "Design e Disponibilidade",
            "Manter um site com design limpo e minimalista, utilizando apenas alertas de evento "
            "para informar os clientes. A disponibilidade contínua é imprescindível; qualquer "
            "interrupção deve ser comunicada com antecedência por alertas visíveis ou outros canais."
        ),
        (
            "Restrições de Conteúdo",
            "É expressamente proibido o uso de links ou a divulgação de conteúdo relacionado a "
            "sites de pornografia, conforme as políticas de uso estabelecidas."
        ),
        (
            "Política de Privacidade e Proteção de Dados",
            "O site deve disponibilizar, de forma clara e acessível, sua Política de Privacidade e "
            "Proteção de Dados, em conformidade com a LGPD (Lei nº 13.709/2018), preferencialmente "
            "no rodapé de todas as páginas."
        ),
    ]

    for titulo, texto in recomendacoes:
        p = doc.add_paragraph()
        _run(p, f"{titulo}: ", bold=True, size=FONT_SMALL)
        _run(p, texto, size=FONT_SMALL)
        set_paragraph_spacing(p, 0, 4)

    doc.add_paragraph()


# ── Conclusão (gerada pelo M3) ────────────────────────────────────────────────

def _add_conclusao(doc: Document, rd: dict) -> None:
    _section_label(doc, "Conclusão")

    overall = rd.get("overall_status", "NÃO VERIFICÁVEL")
    conclusao = rd.get("conclusao", "")

    p = doc.add_paragraph()
    _run(p, conclusao, size=FONT_BODY)
    set_paragraph_spacing(p, 0, 6)

    # Status global em destaque
    p2 = doc.add_paragraph()
    _run(p2, "Status Final: ", bold=True)
    color = GREEN if overall == "CONFORME" else RED
    _run(p2, overall, bold=True, color=color)

    # Data da análise
    p3 = doc.add_paragraph()
    _run(p3, f"Data da análise: {rd.get('analysis_date', date.today().isoformat())}",
         size=FONT_SMALL, color=GRAY)
    set_paragraph_spacing(p3, 4, 0)


# ── Rodapé institucional ──────────────────────────────────────────────────────

def _add_footer_text(doc: Document) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    _run(p, REPORT_FOOTER, size=FONT_SMALL, italic=True, color=GRAY)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, 8, 0)

    p2 = doc.add_paragraph()
    _run(p2, "Documento Restrito", bold=True, size=FONT_SMALL, color=GRAY)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ── Utilitários de formatação ─────────────────────────────────────────────────

def _section_label(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = FONT_NAME
    r.font.size = FONT_SECTION
    r.font.bold = True
    r.font.color.rgb = BLUE_HD
    set_paragraph_spacing(p, before=8, after=4)


def _body_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _run(p, text)
    set_paragraph_spacing(p, 0, 4)


def _run(
    para,
    text: str,
    bold: bool = False,
    italic: bool = False,
    size: Pt | None = None,
    color=None,
) -> None:
    r = para.add_run(text)
    r.font.name = FONT_NAME
    r.font.bold = bold
    r.font.italic = italic
    if size:
        r.font.size = size
    if color:
        r.font.color.rgb = color


def _hdr_cell(cell, text: str) -> None:
    set_cell_bg(cell, "0D47A1")
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.name = FONT_NAME
    r.font.size = FONT_SMALL
    r.font.bold = True
    r.font.color.rgb = WHITE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    remove_paragraph_spacing(p)
    set_cell_vertical_align(cell, "center")


def _body_cell(
    cell,
    text: str,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.name = FONT_NAME
    r.font.size = FONT_SMALL
    r.font.color.rgb = DARK
    p.alignment = align
    remove_paragraph_spacing(p)
    set_cell_vertical_align(cell, "center")


def _status_cell(cell, status: str) -> None:
    p = cell.paragraphs[0]
    r = p.add_run(status)
    r.font.name = FONT_NAME
    r.font.size = FONT_SMALL
    r.font.bold = True
    r.font.color.rgb = STATUS_COLOR.get(status, GRAY)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    remove_paragraph_spacing(p)
    set_cell_vertical_align(cell, "center")
